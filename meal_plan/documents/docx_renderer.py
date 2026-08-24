from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .copy import copy_for, slot_label, day_label
from .helpers import local_food_name, rounded, review_warning_lines
from .models import DocumentContext
from .theme import BORDER, GRAPHITE, INK, IVORY, ORANGE, ORANGE_SOFT, PAPER


FONT_NAME = "Noto Sans Ethiopic"
FALLBACK_FONT = "Ebrima"


def _rgb(hex_value: str) -> RGBColor:
    value = hex_value.lstrip("#")
    return RGBColor.from_string(value.upper())


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def _set_cell_border(cell, *, color: str = BORDER, size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color.lstrip("#"))


def _set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = _rgb(color)


def _style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = _rgb(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 27, INK),
        ("Heading 1", 18, INK),
        ("Heading 2", 13, INK),
        ("Heading 3", 10.5, ORANGE),
    ]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _rgb(color)

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.45)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)


def _orange_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    _set_run_font(run, size=8.5, bold=True, color=ORANGE)


def _section_title(doc: Document, text: str, kicker: str | None = None) -> None:
    if kicker:
        _orange_kicker(doc, kicker)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    _set_run_font(run, size=18, bold=True, color=INK)


def _metric_table(doc: Document, metrics: list[tuple[str, str]]) -> None:
    cols = 3
    rows = (len(metrics) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (label, value) in enumerate(metrics):
        r, c = divmod(idx, cols)
        cell = table.cell(r, c)
        cell.width = Cm(5.5)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(cell, PAPER)
        _set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        _set_run_font(run, size=15, bold=True, color=INK)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(5)
        r2 = p2.add_run(label)
        _set_run_font(r2, size=8.3, bold=False, color=GRAPHITE)
    # blank cells stay subtle
    for idx in range(len(metrics), rows * cols):
        r, c = divmod(idx, cols)
        _shade_cell(table.cell(r, c), IVORY)
        _set_cell_border(table.cell(r, c), color=IVORY)


def _meal_card(doc: Document, meal: dict[str, Any], language: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    _shade_cell(cell, PAPER)
    _set_cell_border(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    slot = slot_label(str(meal.get("slot") or "Meal"), language)
    r = p.add_run(slot.upper())
    _set_run_font(r, size=8.5, bold=True, color=ORANGE)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    r2 = p2.add_run(str(meal.get("meal_name") or "Meal"))
    _set_run_font(r2, size=12, bold=True, color=INK)

    macros = meal.get("macros") or {}
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(5)
    macro_text = (
        f"{rounded(macros.get('kcal'))} kcal   |   "
        f"P {rounded(macros.get('protein'))} g   "
        f"C {rounded(macros.get('carbs'))} g   "
        f"F {rounded(macros.get('fat'))} g"
    )
    r3 = p3.add_run(macro_text)
    _set_run_font(r3, size=8.5, color=GRAPHITE)

    for item in meal.get("items") or []:
        food_name = local_food_name(str(item.get("food_id") or ""), str(item.get("food_name") or "Food"), language)
        grams = rounded(item.get("grams"), 0)
        familiar = str(item.get("familiar") or "").strip()
        line = cell.add_paragraph(style=None)
        line.paragraph_format.left_indent = Cm(0.15)
        line.paragraph_format.space_after = Pt(1)
        bullet = line.add_run("• ")
        _set_run_font(bullet, size=9, bold=True, color=ORANGE)
        name_run = line.add_run(food_name)
        _set_run_font(name_run, size=9, bold=True, color=INK)
        detail = f"  {grams} g"
        if familiar:
            detail += f"  ·  {familiar}"
        detail_run = line.add_run(detail)
        _set_run_font(detail_run, size=8.4, color=GRAPHITE)

    exchanges = meal.get("exchange_options") or []
    if exchanges:
        first = exchanges[0]
        options = first.get("options") or []
        if options:
            opt = options[0]
            p4 = cell.add_paragraph()
            p4.paragraph_format.space_before = Pt(3)
            p4.paragraph_format.space_after = Pt(1)
            c = copy_for(language)
            rr = p4.add_run(c["swap"] + ": ")
            _set_run_font(rr, size=8.3, bold=True, color=ORANGE)
            swap_name = local_food_name(str(opt.get("food_id") or ""), str(opt.get("food_name") or ""), language)
            rr2 = p4.add_run(f"{swap_name} - {rounded(opt.get('exchange_weight_g'))} g")
            _set_run_font(rr2, size=8.3, color=GRAPHITE)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def _cover(doc: Document, plan: dict[str, Any], context: DocumentContext) -> None:
    c = copy_for(context.normalized_language)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("HILAWE")
    _set_run_font(r, size=11, bold=True, color=ORANGE)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(38)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(c["personalized"])
    _set_run_font(r2, size=27, bold=True, color=INK)

    duration = int((plan.get("product") or {}).get("duration_days") or 7)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(24)
    r3 = p3.add_run(f"{duration} DAY  ·  {c['nutrition_system']}")
    _set_run_font(r3, size=10, bold=True, color=GRAPHITE)

    line_table = doc.add_table(rows=1, cols=2)
    line_table.autofit = False
    line_table.columns[0].width = Cm(0.22)
    line_table.columns[1].width = Cm(15.5)
    _shade_cell(line_table.cell(0, 0), ORANGE)
    _shade_cell(line_table.cell(0, 1), ORANGE_SOFT)
    _set_cell_border(line_table.cell(0, 0), color=ORANGE)
    _set_cell_border(line_table.cell(0, 1), color=ORANGE_SOFT)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    rr = p4.add_run(c["prepared_for"] + "\n")
    _set_run_font(rr, size=9, color=GRAPHITE)
    rr2 = p4.add_run(context.client_name)
    _set_run_font(rr2, size=21, bold=True, color=INK)

    p5 = doc.add_paragraph()
    goal = str((plan.get("profile_summary") or {}).get("goal") or "-").replace("_", " ").title()
    rr3 = p5.add_run(goal)
    _set_run_font(rr3, size=11, bold=True, color=ORANGE)

    p6 = doc.add_paragraph()
    p6.paragraph_format.space_before = Pt(42)
    rr4 = p6.add_run(f"{c['plan_id']}: {context.plan_public_id}    ·    {c['version']}: V{context.version_number}")
    _set_run_font(rr4, size=8, color=GRAPHITE)

    if context.status != "APPROVED":
        p7 = doc.add_paragraph()
        rr5 = p7.add_run(c["draft_banner"])
        _set_run_font(rr5, size=8.5, bold=True, color=ORANGE)

    doc.add_page_break()
def _plan_glance(doc: Document, plan: dict[str, Any], context: DocumentContext) -> None:
    c = copy_for(context.normalized_language)
    _section_title(doc, c["plan_glance"], "01")
    targets = plan.get("nutrition_targets") or {}
    product = plan.get("product") or {}
    profile = plan.get("profile_summary") or {}
    client = context.client_profile or {}
    metrics = [
        (c["current_weight"], f"{rounded(client.get('current_weight_kg'), 1)} kg" if client.get("current_weight_kg") else c["not_provided"]),
        (c["target_weight"], f"{rounded(client.get('target_weight_kg'), 1)} kg" if client.get("target_weight_kg") else c["not_provided"]),
        (c["daily_energy"], f"{rounded(targets.get('target_kcal'))} kcal"),
        (c["protein"], f"{rounded(targets.get('protein_g'))} g"),
        (c["meals_day"], str(product.get("meals_per_day") or "-")),
        (c["food_style"], str(profile.get("cuisine_style") or "-").replace("_", " ").title()),
    ]
    _metric_table(doc, metrics)

    doc.add_paragraph()
    summary = [
        (c["goal"], str(profile.get("goal") or "-").replace("_", " ").title()),
        (c["training"], f"{profile.get('training_days_per_week', '-')} days/week · {str(profile.get('training_type') or '-').replace('_', ' ').title()}"),
        (c["budget"], str(profile.get("grocery_budget") or "-").title()),
        (c["diet"], str(profile.get("dietary_pattern") or "-").replace("_", " ").title()),
        (c["fasting"], str(profile.get("orthodox_fasting") or "-").replace("_", " ").title()),
    ]
    table = doc.add_table(rows=len(summary), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(summary):
        left, right = table.rows[i].cells
        _shade_cell(left, IVORY)
        _shade_cell(right, PAPER)
        _set_cell_border(left)
        _set_cell_border(right)
        p = left.paragraphs[0]
        r = p.add_run(label)
        _set_run_font(r, size=8.3, bold=True, color=GRAPHITE)
        p2 = right.paragraphs[0]
        r2 = p2.add_run(value)
        _set_run_font(r2, size=8.8, bold=True, color=INK)
    doc.add_page_break()


def _how_to_use(doc: Document, language: str) -> None:
    c = copy_for(language)
    _section_title(doc, c["how_to_use"], "02")
    for number, title_key, body_key in [
        ("01", "how_1_title", "how_1"),
        ("02", "how_2_title", "how_2"),
        ("03", "how_3_title", "how_3"),
    ]:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(1.3)
        table.columns[1].width = Cm(14.5)
        a, b = table.rows[0].cells
        _shade_cell(a, ORANGE)
        _shade_cell(b, PAPER)
        _set_cell_border(a, color=ORANGE)
        _set_cell_border(b)
        pa = a.paragraphs[0]
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ra = pa.add_run(number)
        _set_run_font(ra, size=10, bold=True, color="#FFFFFF")
        pb = b.paragraphs[0]
        rb = pb.add_run(c[title_key])
        _set_run_font(rb, size=10, bold=True, color=INK)
        pb2 = b.add_paragraph()
        rb2 = pb2.add_run(c[body_key])
        _set_run_font(rb2, size=8.7, color=GRAPHITE)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    doc.add_page_break()


def _rotation(doc: Document, plan: dict[str, Any], language: str) -> None:
    c = copy_for(language)
    _section_title(doc, c["month_map"], "03")
    rotation = plan.get("rotation") or []
    weeks: dict[int, set[str]] = {}
    for row in rotation:
        weeks.setdefault(int(row.get("week") or 1), set()).add(str(row.get("mode") or "PRIMARY"))
    table = doc.add_table(rows=max(1, len(weeks)), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if not weeks:
        weeks = {1: {"PRIMARY"}}
    for i, (week, modes) in enumerate(sorted(weeks.items())):
        a, b, d = table.rows[i].cells
        mode = "SWAP" if "SWAP" in modes else "PRIMARY"
        values = (f"{c['week']} {week}", c["swap_rotation"] if mode == "SWAP" else c["primary"], "")
        for cell, value in zip((a, b, d), values):
            _shade_cell(cell, PAPER if cell is not a else ORANGE_SOFT)
            _set_cell_border(cell)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            _set_run_font(r, size=9, bold=True if cell is not d else False, color=INK if cell is not a else ORANGE)
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"7 days = one personalized core week. 14/30-day products use the same reviewed core with the approved rotation and swaps shown above.")
    _set_run_font(r, size=8.8, color=GRAPHITE)
    fasting_dates = [str(row.get("date")) for row in rotation if row.get("core_source") == "FASTING"]
    if fasting_dates:
        p2 = doc.add_paragraph()
        label = "የወቅታዊ ጾም መዋቅር የሚጠቀሙ ቀናት: " if language == "AM" else "Dates using the seasonal fasting core: "
        _set_run_font(p2.add_run(label + ", ".join(fasting_dates)), size=8.5, bold=True, color=ORANGE)
    doc.add_page_break()


def _days(doc: Document, plan: dict[str, Any], language: str) -> None:
    c = copy_for(language)
    rows = list(plan.get("core_week") or [])
    fasting_rows = list(plan.get("fasting_core_week") or [])
    for day in rows:
        _orange_kicker(doc, f"DAY {int(day.get('day_index', 0)) + 1:02d}  ·  {day.get('date', '')}")
        p = doc.add_paragraph(style="Heading 1")
        r = p.add_run(day_label(str(day.get("day_name") or "Day"), language))
        _set_run_font(r, size=20, bold=True, color=INK)
        if day.get("fasting"):
            rr = p.add_run("   " + c["fasting_day"])
            _set_run_font(rr, size=8.5, bold=True, color=ORANGE)
        totals = day.get("totals") or {}
        p2 = doc.add_paragraph()
        r2 = p2.add_run(
            f"{rounded(totals.get('kcal'))} kcal   |   P {rounded(totals.get('protein'))} g   "
            f"C {rounded(totals.get('carbs'))} g   F {rounded(totals.get('fat'))} g"
        )
        _set_run_font(r2, size=8.8, color=GRAPHITE)
        for meal in day.get("meals") or []:
            _meal_card(doc, meal, language)
        warnings = day.get("warnings") or []
        if warnings:
            p3 = doc.add_paragraph()
            rr = p3.add_run(c["warning"] + ": " + " | ".join(str(x) for x in warnings[:2]))
            _set_run_font(rr, size=7.8, color=ORANGE)
        doc.add_page_break()

    if fasting_rows:
        _section_title(doc, c["fasting_core"], "03F")
        note = doc.add_paragraph()
        _set_run_font(note.add_run(c["fasting_core_note"]), size=8.8, color=GRAPHITE)
        for day in fasting_rows:
            _orange_kicker(doc, f"DAY {int(day.get('day_index', 0)) + 1:02d}  ·  {day.get('date', '')}")
            p = doc.add_paragraph(style="Heading 1")
            _set_run_font(p.add_run(day_label(str(day.get("day_name") or "Day"), language)), size=20, bold=True, color=INK)
            _set_run_font(p.add_run("   " + c["fasting_day"]), size=8.5, bold=True, color=ORANGE)
            totals = day.get("totals") or {}
            p2 = doc.add_paragraph()
            _set_run_font(p2.add_run(
                f"{rounded(totals.get('kcal'))} kcal   |   P {rounded(totals.get('protein'))} g   "
                f"C {rounded(totals.get('carbs'))} g   F {rounded(totals.get('fat'))} g"
            ), size=8.8, color=GRAPHITE)
            for meal in day.get("meals") or []:
                _meal_card(doc, meal, language)
            doc.add_page_break()


def _grocery(doc: Document, plan: dict[str, Any], language: str) -> None:
    c = copy_for(language)
    _section_title(doc, c["grocery"], "04")
    p = doc.add_paragraph()
    r = p.add_run(c["grocery_intro"])
    _set_run_font(r, size=8.8, color=GRAPHITE)
    rows = plan.get("grocery") or []
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Item", "Category", c["planned"], c["buy"]]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _shade_cell(cell, ORANGE)
        p = cell.paragraphs[0]
        rr = p.add_run(text)
        _set_run_font(rr, size=8, bold=True, color="#FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        values = [
            local_food_name(str(row.get("food_id") or ""), str(row.get("buy_item") or ""), language),
            str(row.get("category") or ""),
            f"{rounded(row.get('planned_grams'))} g",
            str(row.get("purchase_quantity") or ""),
        ]
        for i, (cell, value) in enumerate(zip(cells, values)):
            _shade_cell(cell, PAPER if len(table.rows) % 2 else IVORY)
            _set_cell_border(cell)
            p = cell.paragraphs[0]
            rr = p.add_run(value)
            _set_run_font(rr, size=7.6 if i == 0 else 7.2, bold=i == 0, color=INK if i == 0 else GRAPHITE)
    doc.add_page_break()
    fasting_rows = plan.get("fasting_grocery") or []
    if fasting_rows:
        _section_title(doc, c["fasting_grocery"], "04F")
        fasting_table = doc.add_table(rows=1, cols=3)
        fasting_table.style = "Table Grid"
        for i, text in enumerate(("Item", c["planned"], c["buy"])):
            cell = fasting_table.rows[0].cells[i]
            _shade_cell(cell, ORANGE)
            _set_run_font(cell.paragraphs[0].add_run(text), size=8, bold=True, color="#FFFFFF")
        for row in fasting_rows:
            cells = fasting_table.add_row().cells
            values = (
                local_food_name(str(row.get("food_id") or ""), str(row.get("buy_item") or ""), language),
                f"{rounded(row.get('planned_grams'))} g",
                str(row.get("purchase_quantity") or ""),
            )
            for i, (cell, value) in enumerate(zip(cells, values)):
                _set_cell_border(cell)
                _set_run_font(cell.paragraphs[0].add_run(value), size=7.5, bold=i == 0, color=INK)
        doc.add_page_break()


def _guides(doc: Document, context: DocumentContext) -> None:
    c = copy_for(context.normalized_language)
    _section_title(doc, c["portion_hydration"], "05")
    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run(c["hydration"])
    _set_run_font(r, size=13, bold=True, color=INK)
    if context.hydration_target_l:
        p2 = doc.add_paragraph()
        rr = p2.add_run(f"{context.hydration_target_l:.1f} L / day")
        _set_run_font(rr, size=22, bold=True, color=ORANGE)
    p3 = doc.add_paragraph()
    rr3 = p3.add_run(c["hydration_general"])
    _set_run_font(rr3, size=9, color=GRAPHITE)
    doc.add_paragraph()
    p4 = doc.add_paragraph(style="Heading 2")
    rr4 = p4.add_run(c["exact"] + " + " + c["familiar"])
    _set_run_font(rr4, size=13, bold=True, color=INK)
    p5 = doc.add_paragraph()
    rr5 = p5.add_run(c["portion_note"])
    _set_run_font(rr5, size=9, color=GRAPHITE)
    doc.add_paragraph()
    p6 = doc.add_paragraph(style="Heading 2")
    rr6 = p6.add_run(c["coach_note"])
    _set_run_font(rr6, size=13, bold=True, color=ORANGE)
    p7 = doc.add_paragraph()
    rr7 = p7.add_run(c["coach_text"])
    _set_run_font(rr7, size=10, color=INK)
    doc.add_page_break()


def _review(doc: Document, plan: dict[str, Any], context: DocumentContext) -> None:
    c = copy_for(context.normalized_language)
    _section_title(doc, c["review"], "06")
    p = doc.add_paragraph()
    r = p.add_run(c["review_required"])
    _set_run_font(r, size=10, bold=True, color=ORANGE)
    values = [
        (c["plan_id"], context.plan_public_id),
        (c["version"], f"V{context.version_number}"),
        (c["status"], context.status),
        (c["engine"], str(plan.get("engine_version") or "-")),
        (c["dataset"], str(plan.get("dataset_version") or "-")),
    ]
    if context.approved_by:
        values.append((c["approved_by"], context.approved_by))
    if context.approved_at:
        values.append((c["approved_at"], context.approved_at))
    table = doc.add_table(rows=len(values), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(values):
        a, b = table.rows[i].cells
        _shade_cell(a, ORANGE_SOFT)
        _shade_cell(b, PAPER)
        _set_cell_border(a)
        _set_cell_border(b)
        ra = a.paragraphs[0].add_run(label)
        _set_run_font(ra, size=8.2, bold=True, color=ORANGE)
        rb = b.paragraphs[0].add_run(value)
        _set_run_font(rb, size=8.2, color=INK)
    warnings = review_warning_lines(plan)
    if warnings:
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        rr = p2.add_run(c["warning"])
        _set_run_font(rr, size=9, bold=True, color=ORANGE)
        for warning in warnings:
            p3 = doc.add_paragraph(style=None)
            rr3 = p3.add_run("• " + warning)
            _set_run_font(rr3, size=8.3, color=GRAPHITE)


def render_docx(plan: dict[str, Any], context: DocumentContext, output_path: str | Path) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _style_document(doc)
    # Ensure the first section uses compact A4-like margins; default is A4 under most LO/Word installs.
    _cover(doc, plan, context)
    _plan_glance(doc, plan, context)
    _how_to_use(doc, context.normalized_language)
    _rotation(doc, plan, context.normalized_language)
    _days(doc, plan, context.normalized_language)
    _grocery(doc, plan, context.normalized_language)
    _guides(doc, context)
    _review(doc, plan, context)
    doc.save(path)
    return path
